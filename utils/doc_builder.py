"""
Word document builder for Multi-Agent Studio.

Takes the final pipeline state and produces a .docx file as bytes.
The bytes are returned directly to Streamlit's download button —
no file is written to disk.

Handles markdown-style headings (## and ###), bullet points (* or -),
and bold markers (**text**) by converting them to proper Word styles.

Color scheme matches the reference documents:
  Title:   #1E3A5F  dark navy,   20pt bold
  H1:      #1E3A5F  dark navy,   16pt bold
  H2:      #2E75B6  medium blue, 13pt bold
  H3:      #C8860A  amber/gold,  12pt bold
  Body:    #1A1A1A  near-black,  11pt
"""

import io
import re
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Imported at module level to avoid a circular import at runtime.
# utils/ is the foundation layer — modules import from utils, not the reverse.
# This import is guarded so the rest of doc_builder works even if m02 is not present.
try:
    from modules.m02_stock.agents import format_metric_value as _format_metric_value
except ImportError:
    def _format_metric_value(metric: str, value) -> str:
        return str(value) if value is not None else "n/a"


# ── Color palette ─────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1E, 0x3A, 0x5F)   # H1 + Title
BLUE   = RGBColor(0x2E, 0x75, 0xB6)   # H2
AMBER  = RGBColor(0xC8, 0x86, 0x0A)   # H3
GREY   = RGBColor(0x88, 0x88, 0x88)   # meta text / footer
BLACK  = RGBColor(0x1A, 0x1A, 0x1A)   # body text
RULE   = RGBColor(0x2E, 0x75, 0xB6)   # horizontal rule under title
GREEN  = RGBColor(0x2E, 0x7D, 0x32)   # positive evidence items
RED    = RGBColor(0xC6, 0x28, 0x28)   # negative evidence items


def _set_document_styles(doc: Document) -> None:
    """
    Configures base Normal style, margins, and heading styles.
    Heading styles are set on the style objects so every heading
    in the document inherits the right color and size automatically.
    """
    # Base font and margins
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Heading 1 — dark navy, 16pt bold
    h1 = doc.styles["Heading 1"]
    h1.font.name       = "Calibri"
    h1.font.size       = Pt(16)
    h1.font.bold       = True
    h1.font.color.rgb  = NAVY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after  = Pt(6)
    h1.paragraph_format.keep_with_next = True

    # Heading 2 — medium blue, 13pt bold
    h2 = doc.styles["Heading 2"]
    h2.font.name       = "Calibri"
    h2.font.size       = Pt(13)
    h2.font.bold       = True
    h2.font.color.rgb  = BLUE
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after  = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # Heading 3 — amber/gold, 12pt bold
    h3 = doc.styles["Heading 3"]
    h3.font.name       = "Calibri"
    h3.font.size       = Pt(12)
    h3.font.bold       = True
    h3.font.color.rgb  = AMBER
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after  = Pt(3)
    h3.paragraph_format.keep_with_next = True


def _add_horizontal_rule(doc: Document) -> None:
    """Adds a thin blue horizontal rule using a bottom border on an empty paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(10)

    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_bold_runs(paragraph, text: str) -> None:
    """
    Adds text to a paragraph, converting **bold** markers to bold runs.
    Preserves font color of the parent paragraph style.
    """
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def _is_bullet(line: str) -> tuple[bool, str]:
    """
    Returns (True, text) if the line is a bullet point.
    Recognises: '* text', '- text', '• text'
    """
    stripped = line.strip()
    for prefix in ("* ", "- ", "• "):
        if stripped.startswith(prefix):
            return True, stripped[len(prefix):]
    return False, stripped


def _add_markdown_content(doc: Document, text: str) -> None:
    """
    Parses markdown-style text and adds it to the document.

    Converts:
        ## Heading   →  Word Heading 1  (navy)
        ### Heading  →  Word Heading 2  (blue)
        #### Heading →  Word Heading 3  (amber)
        * text / - text  →  List Bullet style
        **bold**     →  bold run
        blank line   →  paragraph spacer
        plain text   →  Normal paragraph
    """
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)

        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)

        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)

        elif stripped == "":
            # Blank line — add a small spacer only if previous was not a heading
            doc.add_paragraph("")

        else:
            is_bul, bul_text = _is_bullet(stripped)
            if is_bul:
                para = doc.add_paragraph(style="List Bullet")
                para.paragraph_format.left_indent  = Inches(0.3)
                para.paragraph_format.space_after  = Pt(3)
                _add_bold_runs(para, bul_text)
            else:
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(6)
                _add_bold_runs(para, stripped)

        i += 1


def build_research_doc(state: dict) -> bytes:
    """
    Builds a formatted Word document from the completed research pipeline state.

    Args:
        state: the final ResearchState dict after all agents have run.

    Returns:
        bytes: the .docx file contents, ready for st.download_button.
    """
    doc = Document()
    _set_document_styles(doc)

    topic      = state.get("topic", "Research Report")
    title      = state.get("title") or topic
    final_text = state.get("final", state.get("draft", "No output generated."))
    sources    = state.get("sources", [])
    model_used = state.get("model_used", "unknown")
    today      = date.today().strftime("%B %d, %Y")

    # ── Title block ──────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after  = Pt(4)
    title_run = title_para.add_run(title)
    title_run.font.name  = "Calibri"
    title_run.font.size  = Pt(22)
    title_run.font.bold  = True
    title_run.font.color.rgb = NAVY

    meta_para = doc.add_paragraph()
    meta_para.paragraph_format.space_after = Pt(2)
    meta_run = meta_para.add_run(f"Research Assistant  |  {today}")
    meta_run.font.size      = Pt(9)
    meta_run.font.color.rgb = GREY

    # Blue rule under the title — same visual device as the reference documents
    _add_horizontal_rule(doc)

    # ── Main content ─────────────────────────────────────────────────────────
    _add_markdown_content(doc, final_text)

    # ── Sources ──────────────────────────────────────────────────────────────
    if sources:
        doc.add_paragraph("")
        doc.add_heading("Sources", level=1)
        for i, url in enumerate(sources, 1):
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(url)
            run.font.size = Pt(9)
            run.font.color.rgb = GREY

    # ── Footer ───────────────────────────────────────────────────────────────
    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        f"Generated by Multi-Agent Studio — Research Assistant  |  Model: {model_used}  |  {today}\n"
        "AI-generated output. Review before use."
    )
    footer_run.font.size      = Pt(8)
    footer_run.font.color.rgb = GREY

    # ── Return as bytes ──────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _bullet_list(doc: Document, items: list, empty_text: str = "No findings returned.") -> None:
    """Adds a List Bullet paragraph per item, or one placeholder line if empty."""
    if not items:
        doc.add_paragraph(empty_text)
        return
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Inches(0.3)
        para.paragraph_format.space_after = Pt(3)
        _add_bold_runs(para, str(item))


def build_stock_research_doc(state: dict) -> bytes:
    """
    Builds the Stock Analyser's equity research note as a Word document.

    Renders directly from the structured fields the Synthesizer returned
    (data_bundle, synthesizer_data, bull_case, bear_case) rather than
    re-parsing the on-screen research_note text — the plain-text template
    uses ALL-CAPS headers and unicode rule lines that do not match the
    ## / ### markdown _add_markdown_content() expects, so both renderers
    read from the same structured source instead of one parsing the other.

    Args:
        state: the final Stock Analyser pipeline state — expects
            data_bundle, synthesizer_data, bull_case, bear_case,
            evidence_summary, rating, confidence, model_used, time_horizon.

    Returns:
        bytes: the .docx file contents, ready for st.download_button.
    """
    doc = Document()
    _set_document_styles(doc)

    db           = state.get("data_bundle", {})
    data         = state.get("synthesizer_data", {})
    ticker       = db.get("ticker", state.get("ticker", ""))
    company_name = db.get("company_name", state.get("company_name", ticker))
    time_horizon = state.get("time_horizon", "")
    confidence   = state.get("confidence", data.get("confidence", ""))
    rating       = state.get("rating", data.get("rating", ""))
    model_used   = state.get("model_used", "unknown")
    today        = date.today().strftime("%B %d, %Y")

    # ── Title block ──────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(4)
    title_run = title_para.add_run(f"{company_name} ({ticker})")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = NAVY

    meta_para = doc.add_paragraph()
    meta_para.paragraph_format.space_after = Pt(2)
    meta_run = meta_para.add_run(
        f"Equity Research Note — Educational Purpose Only  |  {db.get('sector', '')}  |  "
        f"{time_horizon}  |  {today}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = GREY

    _add_horizontal_rule(doc)

    # ── Rating / confidence banner ────────────────────────────────────────────
    rc_para = doc.add_paragraph()
    rc_para.paragraph_format.space_after = Pt(4)
    rc_run = rc_para.add_run(f"RATING: {rating}    CONFIDENCE: {confidence}")
    rc_run.font.size = Pt(13)
    rc_run.font.bold = True
    rc_run.font.color.rgb = BLUE
    conf_para = doc.add_paragraph()
    conf_para.paragraph_format.space_after = Pt(10)
    conf_para.add_run(data.get("confidence_explanation", ""))

    # ── Investment thesis ──────────────────────────────────────────────────────
    doc.add_heading("Investment Thesis", level=1)
    p = doc.add_paragraph()
    _add_bold_runs(p, data.get("investment_thesis", ""))

    doc.add_heading("Fundamentals Summary", level=1)
    _bullet_list(doc, data.get("fundamentals_bullets", []))

    doc.add_heading("Business Quality Summary", level=1)
    _bullet_list(doc, data.get("quality_bullets", []))

    doc.add_heading("Key Risks to the Thesis", level=1)
    _bullet_list(doc, data.get("key_risks", []), empty_text="No risks returned.")

    # ── The debate ─────────────────────────────────────────────────────────────
    doc.add_heading("The Debate", level=1)
    doc.add_heading("Bull Case", level=2)
    p = doc.add_paragraph()
    _add_bold_runs(p, state.get("bull_case", ""))
    doc.add_heading("Bear Case", level=2)
    p = doc.add_paragraph()
    _add_bold_runs(p, state.get("bear_case", ""))
    doc.add_heading("Synthesis", level=2)
    p = doc.add_paragraph()
    _add_bold_runs(p, data.get("debate_synthesis", ""))

    # ── Valuation context ─────────────────────────────────────────────────────
    doc.add_heading("Valuation Context", level=1)
    pe_label = "Trailing" if db.get("pe_used") == "trailing" else "Forward"
    gross_margin = db.get("gross_margin") or 0
    operating_margin = db.get("operating_margin") or 0
    for line in (
        f"Current Price: ${db.get('current_price')}   |   52-Week Range: "
        f"${db.get('fifty_two_week_low')} - ${db.get('fifty_two_week_high')}",
        f"{pe_label} P/E: {db.get('pe_value')}",
        f"Gross Margin: {gross_margin * 100:.1f}%   |   Operating Margin: {operating_margin * 100:.1f}%",
    ):
        doc.add_paragraph(line)

    peers = db.get("peers", [])
    if peers:
        doc.add_heading("Peer Comparison", level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        headers = ["Ticker", "P/E", "Gross Margin", "Op. Margin", "Rev. Growth", "Market Cap"]
        for i, label in enumerate(headers):
            table.rows[0].cells[i].text = label
        for peer in peers:
            row = table.add_row().cells
            row[0].text = peer.get("ticker", "")
            row[1].text = f"{peer['pe']:.1f}" if peer.get("pe") is not None else "n/a"
            row[2].text = f"{peer['gross_margin'] * 100:.1f}%" if peer.get("gross_margin") is not None else "n/a"
            row[3].text = f"{peer['operating_margin'] * 100:.1f}%" if peer.get("operating_margin") is not None else "n/a"
            row[4].text = f"{peer['revenue_growth'] * 100:.1f}%" if peer.get("revenue_growth") is not None else "n/a"
            row[5].text = f"${peer['market_cap'] / 1e9:.1f}B" if peer.get("market_cap") is not None else "n/a"

    dist = db.get("analyst_distribution")
    if dist:
        street = (
            f"{dist['strong_buy'] + dist['buy']} Buy, {dist['hold']} Hold, "
            f"{dist['sell'] + dist['strong_sell']} Sell ({db.get('analyst_count')} analysts)"
        )
    else:
        street = f"Consensus: {db.get('analyst_consensus_key')} ({db.get('analyst_count')} analysts)"

    upside = None
    if db.get("analyst_mean_target") and db.get("current_price"):
        upside = (db["analyst_mean_target"] - db["current_price"]) / db["current_price"] * 100

    doc.add_paragraph(f"Street Consensus: {street}")
    doc.add_paragraph(
        f"Mean Price Target: ${db.get('analyst_mean_target')}   |   Implied Upside/Downside: "
        + (f"{upside:+.1f}%" if upside is not None else "n/a")
    )
    doc.add_paragraph(f"This Analysis vs Street: {data.get('street_comparison', '')}")
    dq_para = doc.add_paragraph()
    dq_run = dq_para.add_run(f"Data Quality Score: {db.get('data_quality_score')} / 100 — {db.get('data_quality_label')}")
    dq_run.bold = True
    for note in db.get("data_quality_breakdown", []):
        doc.add_paragraph(f"• {note.capitalize()}", style="List Bullet")

    # ── Fact Check results — mirrors the on-screen checkpoint gate ───────────
    # Without this section the only trace of a fact-check override was one
    # sentence buried inside confidence_explanation, with no claim/actual detail.
    # A reader who only sees the downloaded doc (not the live app) needs the
    # same claim-by-claim detail the on-screen gate shows before proceeding.
    fc_claims = state.get("fact_check_claims", [])
    if fc_claims:
        format_metric_value = _format_metric_value

        doc.add_heading("Fact Check Results", level=1)
        doc.add_paragraph(state.get("fact_check_summary", ""))
        mismatches = [c for c in fc_claims if c.get("verdict") == "Mismatch"]
        if mismatches:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            for i, label in enumerate(["Metric", "Claimed", "Actual", "Made by"]):
                table.rows[0].cells[i].text = label
            for c in mismatches:
                row = table.add_row().cells
                row[0].text = str(c.get("metric", ""))
                row[1].text = str(c.get("claimed_value", ""))
                row[2].text = format_metric_value(c.get("metric", ""), c.get("true_value"))
                row[3].text = str(c.get("source_agent", ""))
            note_para = doc.add_paragraph()
            note_run = note_para.add_run(
                "A mismatch above was reviewed and proceeded past by the user. "
                "Confidence in this report was capped accordingly."
            )
            note_run.font.color.rgb = AMBER
            note_run.bold = True
        else:
            doc.add_paragraph("All checked claims matched the underlying data.")

    doc.add_heading("Trend Summary", level=1)
    doc.add_paragraph(data.get("trend_summary", ""))

    doc.add_heading("Earnings Quality", level=1)
    doc.add_paragraph(data.get("earnings_quality_summary", ""))

    doc.add_heading("Upcoming Catalysts", level=1)
    catalysts = db.get("catalyst_items", [])
    if catalysts:
        for item in catalysts:
            # No token-cost reason to truncate here — this is the downloaded
            # document itself, not an LLM prompt. Show the full snippet.
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(f"{item.get('title', '')}: {item.get('content', '')}")
    else:
        doc.add_paragraph("No upcoming catalysts identified.")

    doc.add_heading("Macro Context", level=1)
    doc.add_paragraph(db.get("macro_context") or "No sector macro context available.")

    doc.add_heading("Investor Type", level=1)
    doc.add_paragraph(f"This thesis suits a {data.get('investor_type', 'value')} investor with a {time_horizon} horizon.")

    # ── Key evidence ───────────────────────────────────────────────────────────
    doc.add_heading("Key Evidence", level=1)
    evidence = state.get("evidence_summary", [])
    pos_count = sum(1 for e in evidence if e.get("sign") == "+")
    neg_count = sum(1 for e in evidence if e.get("sign") == "-")
    doc.add_paragraph(f"{pos_count} positive signal(s), {neg_count} negative signal(s).")
    for item in evidence:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(f"[{item.get('sign', '?')}] {item.get('text', '')}")
        run.font.color.rgb = GREEN if item.get("sign") == "+" else RED

    doc.add_heading("What This Analysis Cannot Know", level=1)
    for line in (
        "Private management guidance not disclosed publicly",
        "Undisclosed material events (pending litigation, M&A, regulatory actions)",
        "Real-time order flow and institutional positioning",
        "Your personal financial situation, risk tolerance, and investment goals",
        "Tax implications of any transaction",
    ):
        doc.add_paragraph(line, style="List Bullet")

    # ── Disclaimer — bordered table cell, per spec ────────────────────────────
    doc.add_paragraph("")
    disclaimer_table = doc.add_table(rows=1, cols=1)
    disclaimer_table.style = "Table Grid"
    cell = disclaimer_table.rows[0].cells[0]
    p = cell.paragraphs[0]
    header_run = p.add_run("DISCLAIMER\n")
    header_run.bold = True
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = GREY
    body_run = p.add_run(
        "This analysis was produced by an AI system for educational and personal learning "
        "purposes only. It is not investment advice. It does not constitute a recommendation "
        "to buy, sell, or hold any security. Past performance is not indicative of future "
        "results. The analysis relies on publicly available data and is subject to the "
        "limitations described in the \"What This Analysis Cannot Know\" section above. "
        "Do not make investment decisions based on this output."
    )
    body_run.font.size = Pt(9)
    body_run.font.color.rgb = GREY

    # ── Footer ───────────────────────────────────────────────────────────────
    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        f"Generated by Multi-Agent Studio — Stock Analyser  |  Model: {model_used}  |  {today}\n"
        "AI-generated output. Review before use. Not investment advice."
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = GREY

    # ── Return as bytes ──────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
